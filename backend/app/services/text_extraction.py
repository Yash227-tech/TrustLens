"""Document text extraction with smart dispatch.

Per spec §3.2 the Ingestion module "extracts text" from PDFs/images/DOCX.
Tesseract is the spec'd OCR engine; for file types where a structured text
layer already exists we use the faster/exact path:

  DOCX        → python-docx (structured paragraphs + tables)
  Vector PDF  → PyMuPDF text layer
  Image PDF   → Tesseract OCR (fallback when text layer is empty)
  PNG / JPEG  → Tesseract OCR

OCR uses Tesseract 5 with English + Hindi language packs.
"""

from __future__ import annotations

import io
import logging
import os
import re

import fitz  # PyMuPDF
import pytesseract
from docx import Document
from PIL import Image

logger = logging.getLogger(__name__)

# eng + the two most common Indian utility-bill / KYC scripts (Hindi, Gujarati).
# Override via env (e.g. add +tam, +ben) without a rebuild of the language packs
# already baked into the image; revert to "eng" if a doc set is pure English.
OCR_LANGUAGES = os.environ.get("OCR_LANGUAGES", "eng+hin+guj")
PDF_TEXT_MIN_CHARS = 50  # If a PDF text layer yields less than this, treat as image-only.
# Scanner apps (very common in India) stamp a boilerplate text layer on an
# otherwise image-only scan — e.g. "Scanned by CamScanner". That watermark alone
# can exceed PDF_TEXT_MIN_CHARS and trick us into skipping OCR, so the real
# (image) content is never read. Strip these before the min-chars decision.
SCANNER_WATERMARK_RE = re.compile(
    r"scanned (?:by|with|using)\s+\w+|cam ?scanner|tap ?scanner|adobe scan|"
    r"doc ?scanner|office ?lens|kaagaz|clear ?scan",
    re.IGNORECASE)
OCR_DPI = 200
MAX_RESPONSE_CHARS = 10_000  # Truncate extracted text in the API response.

PDF_TYPE = "application/pdf"
DOCX_TYPE = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
IMAGE_TYPES = {"image/png", "image/jpeg"}


def _extract_docx_text(content: bytes) -> str:
    doc = Document(io.BytesIO(content))
    parts: list[str] = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def _extract_pdf_text_layer(content: bytes) -> str:
    with fitz.open(stream=content, filetype="pdf") as pdf:
        return "\n".join(page.get_text() for page in pdf)


# Small phone-photo / thumbnail uploads (e.g. a 265x409 passport image) carry
# print too small for Tesseract to resolve — passport numbers, MRZ, IDs. Upscaling
# before OCR markedly improves recognition of fine print on low-res uploads.
OCR_MIN_DIM = 1000


def _ocr_image(image: Image.Image) -> str:
    image = image.convert("RGB")
    if min(image.size) < OCR_MIN_DIM:
        scale = min(OCR_MIN_DIM / min(image.size), 4.0)
        image = image.resize(
            (int(image.width * scale), int(image.height * scale)),
            Image.Resampling.LANCZOS,
        )
    return pytesseract.image_to_string(image, lang=OCR_LANGUAGES)


def _ocr_pdf_pages(content: bytes) -> str:
    parts: list[str] = []
    with fitz.open(stream=content, filetype="pdf") as pdf:
        mat = fitz.Matrix(OCR_DPI / 72, OCR_DPI / 72)
        for page in pdf:
            pix = page.get_pixmap(matrix=mat, alpha=False)
            img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)
            parts.append(pytesseract.image_to_string(img, lang=OCR_LANGUAGES))
    return "\n".join(parts)


def extract_text(content: bytes, content_type: str) -> tuple[str, str]:
    """Extract text from a document. Returns (text, method_label)."""
    if content_type == DOCX_TYPE:
        try:
            return _extract_docx_text(content), "python-docx"
        except Exception as e:
            logger.warning("DOCX text extraction failed: %s", e)
            return "", "docx_failed"

    if content_type == PDF_TYPE:
        try:
            text = _extract_pdf_text_layer(content)
            # Ignore scanner-app watermarks when judging whether a real text layer
            # exists — otherwise a "Scanned by CamScanner" stamp masks an image-only
            # scan and we never OCR the actual content.
            meaningful = SCANNER_WATERMARK_RE.sub("", text).strip()
            if len(meaningful) >= PDF_TEXT_MIN_CHARS:
                return text, "pymupdf-text-layer"
            text = _ocr_pdf_pages(content)
            return text, "tesseract-pdf-ocr"
        except Exception as e:
            logger.warning("PDF text extraction failed: %s", e)
            return "", "pdf_failed"

    if content_type in IMAGE_TYPES:
        try:
            img = Image.open(io.BytesIO(content))
            return _ocr_image(img), "tesseract-image-ocr"
        except Exception as e:
            logger.warning("Image OCR failed: %s", e)
            return "", "image_failed"

    return "", "unsupported"


def truncate_for_response(text: str) -> str:
    if len(text) <= MAX_RESPONSE_CHARS:
        return text
    return text[:MAX_RESPONSE_CHARS] + "\n… [truncated, full text retained internally]"
